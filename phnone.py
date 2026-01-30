import streamlit as st
from docx import Document
import io  # メモリ上でファイルを扱うための道具

# --- 1. 画面のレイアウトを作る ---
st.title("📄 スマホでWord作成アプリ")
st.write("以下の項目を入力して、ボタンを押してください。")

# ユーザー入力欄（テキスト入力）
user_name = st.text_input("お名前", "山田 太郎")
report_content = st.text_area("報告内容", "ここに今週の業務報告を入力してください。")

# --- 2. Wordを作る機能 ---
def create_word_file(name, text):
    doc = Document()
    # タイトル
    doc.add_heading('業務報告書', 0)
    # 名前
    doc.add_paragraph(f'作成者：{name}')
    # 本文
    doc.add_heading('【内容】', level=1)
    doc.add_paragraph(text)
    
    # ★重要：Webアプリでは「HDD」ではなく「メモリ」に保存する
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue() # ファイルの中身データを返す

# --- 3. ダウンロードボタンの表示 ---
# ドキュメントを作成する中身のデータを用意
word_data = create_word_file(user_name, report_content)

st.write("---") # 区切り線
st.write("準備ができたらダウンロードボタンを押してください👇")

# ダウンロードボタン
st.download_button(
    label="📥 Wordファイルをダウンロード",
    data=word_data,
    file_name="my_report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)